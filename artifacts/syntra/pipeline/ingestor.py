"""Ingestor — bytes → Document (DOCX via python-docx; PDF via LandingAI ADE or pdfplumber)."""
from __future__ import annotations
import io
import os
import re
from models import Document, Element

# Patterns for detecting the start of a new legal section / paragraph break
_NUMBERED_SECTION = re.compile(
    r"^\s*"
    r"(?:"
    r"\d{1,2}(?:\.\d{1,2})*\."   # 1.  /  1.1.  /  2.3.
    r"|[A-Z]{1,5}\."              # A.  /  XIV.
    r"|\([a-z]\)"                  # (a)
    r"|\([ivx]+\)"                 # (iv)
    r")"
    r"\s+[A-Z(\"']",              # followed by capital / open-paren / quote
    re.VERBOSE,
)

_ALL_CAPS_HEADING = re.compile(r"^[A-Z][A-Z\s\-,./':&]{5,}$")


class Ingestor:
    def ingest(self, file_bytes: bytes, source_type: str, doc_id: str) -> Document:
        if source_type == "docx":
            return self._parse_docx(file_bytes, doc_id)
        elif source_type == "pdf":
            return self._parse_pdf(file_bytes, doc_id)
        raise ValueError(f"Unsupported source type: {source_type!r}")

    # ── DOCX ─────────────────────────────────────────────────────────────────

    def _parse_docx(self, file_bytes: bytes, doc_id: str) -> Document:
        from docx import Document as DocxDoc

        docx = DocxDoc(io.BytesIO(file_bytes))
        elements: list[Element] = []
        text_parts: list[str] = []
        offset = 0
        heading_stack: list[str] = []

        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            kind = "paragraph"
            heading_path = list(heading_stack)

            if para.style.name.startswith("Heading"):
                kind = "heading"
                try:
                    level = int(para.style.name.split()[-1])
                except (IndexError, ValueError):
                    level = 1
                heading_stack = heading_stack[: level - 1] + [text]
                heading_path = list(heading_stack)

            elements.append(Element(
                kind=kind, text=text,
                start=offset, end=offset + len(text),
                heading_path=heading_path,
            ))
            text_parts.append(text)
            offset += len(text) + 1

        full_text = "\n".join(text_parts)
        return Document(doc_id=doc_id, source_type="docx", full_text=full_text, elements=elements)

    # ── PDF ──────────────────────────────────────────────────────────────────

    def _parse_pdf(self, file_bytes: bytes, doc_id: str) -> Document:
        api_key = os.environ.get("LANDING_AI_API_KEY")
        if api_key:
            result = self._parse_pdf_ade(file_bytes, doc_id, api_key)
            if result is not None:
                return result
        return self._parse_pdf_plumber(file_bytes, doc_id)

    def _parse_pdf_ade(self, file_bytes: bytes, doc_id: str, api_key: str) -> Document | None:
        try:
            # agentic-doc reads VISION_AGENT_API_KEY from env
            os.environ.setdefault("VISION_AGENT_API_KEY", api_key)

            from agentic_doc.parse import parse  # type: ignore

            # parse() accepts raw bytes directly → returns List[ParsedDocument]
            results = parse(file_bytes)
            if not results:
                return None

            elements: list[Element] = []
            text_parts: list[str] = []
            offset = 0

            for parsed_doc in results:
                for chunk in (parsed_doc.chunks or []):
                    text = (chunk.text or "").strip()
                    if not text:
                        continue
                    # chunk_type values: "text", "table", "heading", "figure", etc.
                    chunk_type = getattr(chunk, "chunk_type", "") or ""
                    kind = "heading" if "heading" in chunk_type.lower() or "title" in chunk_type.lower() else "paragraph"
                    elements.append(Element(
                        kind=kind, text=text,
                        start=offset, end=offset + len(text),
                    ))
                    text_parts.append(text)
                    offset += len(text) + 1

            if not elements:
                return None

            return Document(
                doc_id=doc_id, source_type="pdf",
                full_text="\n".join(text_parts), elements=elements,
            )
        except ImportError:
            return None
        except Exception as exc:
            print(f"[ingestor] LandingAI ADE error: {exc}")
            return None

    def _parse_pdf_plumber(self, file_bytes: bytes, doc_id: str) -> Document:
        """
        Extract text from PDF and join wrapped lines into logical paragraphs.

        A new paragraph starts when:
          - A blank line appears
          - A numbered section starts  (e.g. "3. Remedies.")
          - An all-caps heading appears
          - Indentation changes sharply (heuristic: line shorter than 40 chars
            followed by a non-indented full line suggests a section break)
        """
        try:
            import pdfplumber

            all_lines: list[str] = []
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    raw = page.extract_text() or ""
                    all_lines.extend(raw.split("\n"))

        except ImportError:
            placeholder = "[PDF — install pdfplumber or set LANDING_AI_API_KEY]"
            return Document(
                doc_id=doc_id, source_type="pdf",
                full_text=placeholder,
                elements=[Element(kind="paragraph", text=placeholder,
                                  start=0, end=len(placeholder))],
            )

        # ── Join lines into logical paragraphs ───────────────────────────────
        paragraphs: list[tuple[str, str]] = []  # (kind, text)
        current_lines: list[str] = []

        def flush(kind: str = "paragraph"):
            if current_lines:
                joined = " ".join(current_lines).strip()
                if len(joined) >= 5:
                    paragraphs.append((kind, joined))
                current_lines.clear()

        for raw_line in all_lines:
            stripped = raw_line.strip()

            # Blank line → paragraph break
            if not stripped:
                flush()
                continue

            # All-caps heading
            if _ALL_CAPS_HEADING.match(stripped):
                flush()
                current_lines.append(stripped)
                flush("heading")
                continue

            # Numbered section start → new paragraph
            if _NUMBERED_SECTION.match(stripped):
                flush()
                current_lines.append(stripped)
                continue

            # Continuation of current paragraph
            current_lines.append(stripped)

        flush()

        # ── Build Document ────────────────────────────────────────────────────
        elements: list[Element] = []
        text_parts: list[str] = []
        offset = 0

        for kind, text in paragraphs:
            elements.append(Element(
                kind=kind, text=text,
                start=offset, end=offset + len(text),
            ))
            text_parts.append(text)
            offset += len(text) + 1

        return Document(
            doc_id=doc_id, source_type="pdf",
            full_text="\n".join(text_parts), elements=elements,
        )
