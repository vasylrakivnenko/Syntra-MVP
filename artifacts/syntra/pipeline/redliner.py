"""Redliner — produce a .docx with colour-coded redline summary section (§9)."""
from __future__ import annotations
import io
from models import Document, ClauseVerdict, RuleVerdict


_STATUS_COLOR = {
    "unacceptable":        (0xDC, 0x35, 0x45),   # red
    "acceptable_deviation": (0xFF, 0x8C, 0x00),   # orange
    "unusual":             (0x6F, 0x42, 0xC1),    # purple
    "complies":            (0x19, 0x87, 0x54),    # green
}

_RULE_VERDICT_LABEL = {
    "breach":            ("POSITION BREACHED", (0xDC, 0x35, 0x45)),
    "not_covered":       ("PLAYBOOK GAP", (0xDC, 0x35, 0x45)),
    "attorney_question": ("ATTORNEY QUESTION", (0x6F, 0x42, 0xC1)),
    "met_via_fallback":  ("MET VIA FALLBACK", (0xFF, 0x8C, 0x00)),
}


class Redliner:
    def redline(self, doc: Document, verdicts: list[ClauseVerdict],
                clauses: list | None = None,
                rule_verdicts: list[RuleVerdict] | None = None) -> bytes:
        from docx import Document as DocxDoc
        from docx.shared import RGBColor, Pt

        # clause_id -> clause, so citations can quote the counterparty text.
        clause_map = {c.id: c for c in (clauses or [])}

        rdoc = DocxDoc()

        # ── reproduce original document ──────────────────────────────────────
        for el in doc.elements:
            if el.kind == "heading" and el.heading_path:
                level = min(len(el.heading_path), 9)
                rdoc.add_heading(el.text, level=level)
            else:
                rdoc.add_paragraph(el.text)

        # ── redline summary section ──────────────────────────────────────────
        rdoc.add_page_break()
        rdoc.add_heading("REDLINE SUMMARY — ISSUES FLAGGED BY SYNTRA", level=1)

        # Rule-level findings first: one row per playbook position, reconciled
        # across the whole document. Per-clause flags below remain the evidence.
        rule_rows = [r for r in (rule_verdicts or [])
                     if r.verdict in _RULE_VERDICT_LABEL]
        if rule_rows:
            rdoc.add_heading("Findings by playbook position", level=2)
            for r in rule_rows:
                label, rgb = _RULE_VERDICT_LABEL[r.verdict]
                p = rdoc.add_paragraph()
                run = p.add_run(f"[{label}] ")
                run.bold = True
                run.font.color.rgb = RGBColor(*rgb)
                name = ((r.cited_position or {}).get("policy_label")
                        or (r.clause_type or "position").replace("_", " ").title())
                p.add_run(name).bold = True
                if r.rationale:
                    rp = rdoc.add_paragraph()
                    rp.add_run(r.rationale).italic = True
                if r.question:
                    qp = rdoc.add_paragraph()
                    qr = qp.add_run(f"QUESTION FOR ATTORNEY: {r.question}")
                    qr.font.color.rgb = RGBColor(0x6F, 0x42, 0xC1)
                if r.suggested_text:
                    sp = rdoc.add_paragraph()
                    sr = sp.add_run(f"SUGGESTION: {r.suggested_text}")
                    sr.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
                rdoc.add_paragraph()
            rdoc.add_heading("Per-clause evidence", level=2)

        action_verdicts = [
            v for v in verdicts
            if v.branch == "verdict" and v.status != "complies"
        ]
        silences = [v for v in verdicts if v.branch == "silence"]

        if not action_verdicts and not silences:
            p = rdoc.add_paragraph()
            p.add_run("✓ No material issues found.").bold = True
            return self._to_bytes(rdoc)

        # Missing-clause flags (silence)
        if silences:
            rdoc.add_heading("Missing required clauses", level=2)
            for v in silences:
                p = rdoc.add_paragraph(style="List Bullet")
                run = p.add_run(f"[SILENCE] {v.reason or 'Required clause absent'}")
                run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
                run.bold = True
                self._add_citation(rdoc, v, clause_map)

        # Clause-level issues
        if action_verdicts:
            rdoc.add_heading("Clause-level issues", level=2)

        for v in action_verdicts:
            rgb = _STATUS_COLOR.get(v.status or "unusual", (0x6C, 0x75, 0x7D))
            color = RGBColor(*rgb)

            # Status + rule badge (rule id shown here only when there is no
            # citation block, which already names the rule)
            p = rdoc.add_paragraph()
            tag_run = p.add_run(f"[{(v.status or '').upper().replace('_', ' ')}] ")
            tag_run.bold = True
            tag_run.font.color.rgb = color
            if v.rule_ids and not v.cited_position:
                rule_run = p.add_run(f"Rule: {', '.join(v.rule_ids)}")
                rule_run.italic = True
                rule_run.font.size = Pt(9)

            # Rationale
            if v.rationale:
                p2 = rdoc.add_paragraph()
                p2.add_run(v.rationale).italic = True

            # Suggested alternative
            if v.suggested_text:
                p3 = rdoc.add_paragraph()
                sugg = p3.add_run(f"SUGGESTION: {v.suggested_text}")
                sugg.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)

            # Grounded citations: the playbook position this finding was judged
            # against and the counterparty clause it came from.
            self._add_citation(rdoc, v, clause_map)

            # Separator
            rdoc.add_paragraph()

        return self._to_bytes(rdoc)

    @staticmethod
    def _add_citation(rdoc, v: ClauseVerdict, clause_map: dict) -> None:
        from docx.shared import RGBColor, Pt

        lines: list[str] = []
        cp = v.cited_position or {}
        if cp:
            label = cp.get("rule_id") or cp.get("policy_label") or "playbook position"
            version = cp.get("playbook_version")
            basis = f"playbook v{version}, as of analysis" if version else "playbook"
            lines.append(f"Company position: {label} — "
                         f"{cp.get('policy_label') or cp.get('clause_type') or ''} ({basis})")
            if cp.get("required") and not cp.get("rule_id"):
                lines.append(f"  Required policy \"{cp.get('policy_label')}\" has no "
                             f"position defined for service line {cp.get('service_line')}.")
            else:
                for key, name in (("preferred", "Preferred"), ("fallback", "Fallback"),
                                  ("walk_away", "Walk-away")):
                    if cp.get(key):
                        lines.append(f"  {name}: \"{cp[key]}\"")

        clause = clause_map.get(v.clause_id)
        if clause is not None:
            heading = " > ".join(clause.heading_path or [])
            excerpt = " ".join((clause.text or "").split())
            if len(excerpt) > 220:
                excerpt = excerpt[:220].rstrip() + "…"
            where = f" ({heading})" if heading else ""
            lines.append(f"Counterparty clause{where}: \"{excerpt}\"")

        if not lines:
            return
        gray = RGBColor(0x6C, 0x75, 0x7D)
        hdr = rdoc.add_paragraph()
        hr = hdr.add_run("SOURCES")
        hr.bold = True
        hr.font.size = Pt(8)
        hr.font.color.rgb = gray
        for line in lines:
            lp = rdoc.add_paragraph()
            lr = lp.add_run(line)
            lr.font.size = Pt(8.5)
            lr.font.color.rgb = gray

    @staticmethod
    def _to_bytes(rdoc) -> bytes:
        buf = io.BytesIO()
        rdoc.save(buf)
        return buf.getvalue()
